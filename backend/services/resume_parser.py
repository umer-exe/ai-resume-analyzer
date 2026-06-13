"""Extract plain text from supported resume files."""

from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_BYTES = 2 * 1024 * 1024
MAX_FILE_SIZE_MESSAGE = "Resume file must be 2 MB or smaller"


class ResumeParseError(ValueError):
    """Raised when an uploaded resume cannot be safely parsed."""


class ResumeFileTooLargeError(ResumeParseError):
    """Raised when an uploaded resume exceeds the file-size limit."""


def _normalize_text(text):
    normalized_lines = (
        text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    )
    return "\n".join(
        line.strip() for line in normalized_lines if line.strip()
    )


def _extract_pdf(file_bytes):
    try:
        with pymupdf.open(stream=file_bytes, filetype="pdf") as document:
            if document.needs_pass:
                raise ResumeParseError(
                    "Password-protected PDF files are not supported"
                )
            return "\n".join(page.get_text() for page in document)
    except ResumeParseError:
        raise
    except Exception as error:
        raise ResumeParseError("Unable to read the PDF file") from error


def _extract_docx(file_bytes):
    try:
        document = Document(BytesIO(file_bytes))
        text_parts = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)
    except Exception as error:
        raise ResumeParseError("Unable to read the DOCX file") from error


def _extract_txt(file_bytes):
    try:
        return file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise ResumeParseError(
            "TXT files must use UTF-8 text encoding"
        ) from error


def extract_resume_text(uploaded_file):
    """Return normalized resume text from a Werkzeug uploaded file."""

    filename = uploaded_file.filename or ""
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParseError(
            "Unsupported file type. Upload a PDF, DOCX, or TXT file"
        )

    file_bytes = uploaded_file.read()
    if not file_bytes:
        raise ResumeParseError("Uploaded file is empty")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ResumeFileTooLargeError(MAX_FILE_SIZE_MESSAGE)

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
    }
    extracted_text = _normalize_text(extractors[extension](file_bytes))

    if not extracted_text:
        raise ResumeParseError("No readable text found in resume")

    return extracted_text
