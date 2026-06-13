import unittest
from io import BytesIO
from unittest.mock import patch

import pymupdf
from docx import Document
from werkzeug.datastructures import FileStorage

from app import app
from services.resume_parser import ResumeParseError, extract_resume_text


SAMPLE_TEXT = "Python developer\nFlask and SQL projects"
ANALYSIS_RESULT = {
    "score": 50,
    "status": "Moderate",
    "summary": "Test analysis",
    "checks": {"passed": 2, "warnings": 2, "issues": 2},
    "category_analysis": [],
    "skills": ["Python", "Flask", "SQL"],
    "action_plan": [],
    "ml_prediction": {
        "predicted_category": None,
        "display_category": None,
        "confidence": 0,
        "source": "ml_classifier",
        "message": "ML model not trained yet",
        "top_predictions": [],
    },
}


def uploaded_file(filename, file_bytes):
    return FileStorage(stream=BytesIO(file_bytes), filename=filename)


def create_pdf_bytes(text=None, password=None):
    document = pymupdf.open()
    page = document.new_page()
    if text:
        page.insert_text((72, 72), text)

    save_options = {}
    if password:
        save_options = {
            "encryption": pymupdf.PDF_ENCRYPT_AES_256,
            "owner_pw": "owner-password",
            "user_pw": password,
        }

    file_bytes = document.tobytes(**save_options)
    document.close()
    return file_bytes


def create_docx_bytes(paragraph_text=None, table_text=None):
    document = Document()
    if paragraph_text:
        document.add_paragraph(paragraph_text)
    if table_text:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = table_text

    output = BytesIO()
    document.save(output)
    return output.getvalue()


class ResumeParserTests(unittest.TestCase):
    def test_extracts_utf8_txt_with_optional_bom(self):
        result = extract_resume_text(
            uploaded_file("resume.TXT", b"\xef\xbb\xbfPython\r\n\r\nFlask")
        )

        self.assertEqual(result, "Python\nFlask")

    def test_extracts_pdf_text(self):
        result = extract_resume_text(
            uploaded_file("resume.pdf", create_pdf_bytes("Python developer"))
        )

        self.assertEqual(result, "Python developer")

    def test_extracts_docx_paragraphs_and_table_cells(self):
        result = extract_resume_text(
            uploaded_file(
                "resume.docx",
                create_docx_bytes("Python developer", "Flask project"),
            )
        )

        self.assertEqual(result, "Python developer\nFlask project")

    def test_rejects_unsupported_empty_corrupt_and_textless_files(self):
        invalid_files = [
            ("resume.doc", b"old document"),
            ("resume.txt", b""),
            ("resume.pdf", b"not a PDF"),
            ("resume.pdf", create_pdf_bytes()),
            ("resume.docx", create_docx_bytes()),
        ]

        for filename, file_bytes in invalid_files:
            with self.subTest(filename=filename, size=len(file_bytes)):
                with self.assertRaises(ResumeParseError):
                    extract_resume_text(uploaded_file(filename, file_bytes))

    def test_rejects_password_protected_pdf(self):
        with self.assertRaisesRegex(
            ResumeParseError,
            "Password-protected PDF files are not supported",
        ):
            extract_resume_text(
                uploaded_file(
                    "protected.pdf",
                    create_pdf_bytes("Private resume", password="secret"),
                )
            )


class ResumeUploadApiTests(unittest.TestCase):
    def setUp(self):
        app.config.update(TESTING=True)
        self.client = app.test_client()

    def test_text_submission_keeps_existing_response_schema(self):
        with patch("app.run_analysis", return_value=ANALYSIS_RESULT) as analyzer:
            response = self.client.post(
                "/api/v1/analyze",
                data={"profile_text": SAMPLE_TEXT},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.get_json(),
            {"success": True, "data": ANALYSIS_RESULT, "error": None},
        )
        analyzer.assert_called_once_with(SAMPLE_TEXT)

    def test_supported_uploads_are_extracted_and_analyzed(self):
        uploads = [
            ("resume.txt", SAMPLE_TEXT.encode("utf-8"), SAMPLE_TEXT),
            (
                "resume.pdf",
                create_pdf_bytes("Python developer"),
                "Python developer",
            ),
            (
                "resume.docx",
                create_docx_bytes("Python developer", "Flask project"),
                "Python developer\nFlask project",
            ),
        ]

        for filename, file_bytes, expected_text in uploads:
            with self.subTest(filename=filename):
                with patch(
                    "app.run_analysis",
                    return_value=ANALYSIS_RESULT,
                ) as analyzer:
                    response = self.client.post(
                        "/api/v1/analyze",
                        data={
                            "resume_file": (
                                BytesIO(file_bytes),
                                filename,
                            )
                        },
                    )

                self.assertEqual(response.status_code, 200)
                self.assertEqual(
                    response.get_json(),
                    {
                        "success": True,
                        "data": ANALYSIS_RESULT,
                        "error": None,
                    },
                )
                analyzer.assert_called_once_with(expected_text)

    def test_missing_and_conflicting_inputs_return_400(self):
        missing_response = self.client.post("/api/v1/analyze", data={})
        conflicting_response = self.client.post(
            "/api/v1/analyze",
            data={
                "profile_text": SAMPLE_TEXT,
                "resume_file": (BytesIO(b"Python"), "resume.txt"),
            },
        )

        self.assertEqual(missing_response.status_code, 400)
        self.assertEqual(
            missing_response.get_json()["error"],
            "Profile text or a resume file is required",
        )
        self.assertEqual(conflicting_response.status_code, 400)
        self.assertEqual(
            conflicting_response.get_json()["error"],
            "Provide either profile text or a resume file, not both",
        )

    def test_invalid_uploads_return_400(self):
        invalid_uploads = [
            ("resume.exe", b"not supported"),
            ("resume.txt", b""),
            ("resume.pdf", b"not a PDF"),
            ("resume.pdf", create_pdf_bytes()),
        ]

        for filename, file_bytes in invalid_uploads:
            with self.subTest(filename=filename, size=len(file_bytes)):
                response = self.client.post(
                    "/api/v1/analyze",
                    data={
                        "resume_file": (
                            BytesIO(file_bytes),
                            filename,
                        )
                    },
                )

                self.assertEqual(response.status_code, 400)
                self.assertFalse(response.get_json()["success"])

    def test_oversized_upload_returns_413(self):
        response = self.client.post(
            "/api/v1/analyze",
            data={
                "resume_file": (
                    BytesIO(b"x" * (5 * 1024 * 1024 + 1)),
                    "resume.txt",
                )
            },
        )

        self.assertEqual(response.status_code, 413)
        self.assertEqual(
            response.get_json(),
            {
                "success": False,
                "data": None,
                "error": "Resume file must be 5 MB or smaller",
            },
        )


if __name__ == "__main__":
    unittest.main()
